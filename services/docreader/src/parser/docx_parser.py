import logging
import tempfile
import os
import sys
import time
from io import BytesIO
from typing import Optional
from PIL import Image
from docx import Document
from docx.image.exceptions import (
    UnrecognizedImageError,
    UnexpectedEndOfFileError,
    InvalidImageStreamError,
)
from concurrent.futures import ThreadPoolExecutor, ProcessPoolExecutor, as_completed
import threading
import traceback
from multiprocessing import Manager
import re

from .base_parser import BaseParser

logger = logging.getLogger(__name__)
logger.setLevel(logging.INFO)
# Add thread local storage to track the processing status of each thread
thread_local = threading.local()


class DocxParser(BaseParser):
    """DOCX document parser"""

    def __init__(
        self,
        file_name: str = "",
        file_type: str = None,
        enable_multimodal: bool = True,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        separators: list = ["\n\n", "\n", "。"],
        ocr_backend: str = "paddle",
        ocr_config: dict = None,
        max_image_size: int = 1920,
        max_concurrent_tasks: int = 5,
        max_pages: int = 100,  # Maximum number of pages to process, default to 50 pages
    ):
        """Initialize DOCX document parser

        Args:
            file_name: File name
            file_type: File type, if None, infer from file name
            enable_multimodal: Whether to enable multimodal processing
            chunk_size: Chunk size
            chunk_overlap: Chunk overlap
            separators: List of separators
            ocr_backend: OCR engine type
            ocr_config: OCR engine configuration
            max_image_size: Maximum image size limit
            max_concurrent_tasks: Maximum number of concurrent tasks
            max_pages: Maximum number of pages to process, if more than this, only process the first max_pages pages
        """
        super().__init__(
            file_name=file_name,
            file_type=file_type,
            enable_multimodal=enable_multimodal,
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            separators=separators,
            ocr_backend=ocr_backend,
            ocr_config=ocr_config,
            max_image_size=max_image_size,
            max_concurrent_tasks=max_concurrent_tasks,
        )
        self.max_pages = max_pages
        logger.info(f"DocxParser initialized with max_pages={max_pages}")

    def parse_into_text(self, content: bytes) -> str:
        """Parse DOCX document, only extract text content and image Markdown links

        Args:
            content: DOCX document content

        Returns:
            Parsed text
        """
        logger.info(f"Parsing DOCX document, content size: {len(content)} bytes")
        logger.info(f"Max pages limit set to: {self.max_pages}")
        logger.info("Converting DOCX content to sections and tables")

        start_time = time.time()
        # Use concurrent processing to handle the document
        max_workers = min(4, os.cpu_count() or 2)  # Reduce thread count to avoid excessive memory consumption
        logger.info(f"Setting max_workers to {max_workers} for document processing")

        try:
            logger.info(f"Starting Docx processing with max_pages={self.max_pages}")
            docx_processor = Docx(max_image_size=self.max_image_size,
                                   enable_multimodal=self.enable_multimodal,
                                   upload_file=self.upload_file)
            all_lines, tables = docx_processor(
                binary=content,
                max_workers=max_workers,
                to_page=self.max_pages,
            )
            processing_time = time.time() - start_time
            logger.info(
                f"Docx processing completed in {processing_time:.2f}s, "
                f"extracted {len(all_lines)} sections and {len(tables)} tables"
            )

            # Add debug log to print the structure of all_lines
            if all_lines:
                logger.info(f"all_lines type: {type(all_lines)}")
                if len(all_lines) > 0:
                    first_line = all_lines[0]
                    logger.info(f"First line type: {type(first_line)}, length: {len(first_line)}")
                    logger.info(f"First line content types: {[type(item) for item in first_line]}")
                    if len(all_lines) > 1:
                        second_line = all_lines[1]
                        logger.info(f"Second line type: {type(second_line)}, length: {len(second_line)}")
                        logger.info(f"Second line content types: {[type(item) for item in second_line]}")
            else:
                logger.warning("all_lines is empty")

            logger.info("Processing document sections")
            section_start_time = time.time()

            text_parts = []

            for sec_idx, line in enumerate(all_lines):
                try:
                    text = None
                    if isinstance(line, (tuple, list)) and len(line) >= 1:
                        text = line[0]

                    if text:
                        text_parts.append(text)
                        if sec_idx < 3 or sec_idx % 50 == 0:
                            logger.info(
                                f"Added section {sec_idx+1} text: {text[:50]}..."
                                if len(text) > 50
                                else f"Added section {sec_idx+1} text: {text}"
                            )

                except Exception as e:
                    logger.error(f"Error processing section {sec_idx+1}: {str(e)}")
                    logger.error(f"Detailed stack trace: {traceback.format_exc()}")
                    continue

            # Combine text
            section_processing_time = time.time() - section_start_time
            logger.info(f"Section processing completed in {section_processing_time:.2f}s")
            logger.info("Combining all text parts")
            text = "\n\n".join([part for part in text_parts if part])

            # Check if the generated text is empty
            if not text:
                logger.warning("Generated text is empty, trying alternative method")
                return self._parse_using_simple_method(content)

            total_processing_time = time.time() - start_time
            logger.info(
                f"Parsing complete in {total_processing_time:.2f}s, generated {len(text)} characters of text "
            )

            return text
        except Exception as e:
            logger.error(f"Error parsing DOCX document: {str(e)}")
            logger.error(f"Detailed stack trace: {traceback.format_exc()}")
            return self._parse_using_simple_method(content)

    def _parse_using_simple_method(self, content: bytes) -> str:
        """Parse document using a simplified method, as a fallback

        Args:
            content: Document content

        Returns:
            Parsed text
        """
        logger.info("Attempting to parse document using simplified method")
        start_time = time.time()
        try:
            doc = Document(BytesIO(content))
            logger.info(
                f"Successfully loaded document in simplified method, "
                f"contains {len(doc.paragraphs)} paragraphs and {len(doc.tables)} tables"
            )
            text_parts = []

            # Extract paragraph text
            para_count = len(doc.paragraphs)
            logger.info(f"Extracting text from {para_count} paragraphs")
            para_with_text = 0
            for i, para in enumerate(doc.paragraphs):
                if i % 100 == 0:
                    logger.info(f"Processing paragraph {i+1}/{para_count}")
                if para.text.strip():
                    text_parts.append(para.text.strip())
                    para_with_text += 1

            logger.info(f"Extracted text from {para_with_text}/{para_count} paragraphs")

            # Extract table text
            table_count = len(doc.tables)
            logger.info(f"Extracting text from {table_count} tables")
            tables_with_content = 0
            rows_processed = 0
            for i, table in enumerate(doc.tables):
                if i % 10 == 0:
                    logger.info(f"Processing table {i+1}/{table_count}")

                table_has_content = False
                for row in table.rows:
                    rows_processed += 1
                    row_text = ' | '.join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        text_parts.append(row_text)
                        table_has_content = True

                if table_has_content:
                    tables_with_content += 1

            logger.info(
                f"Extracted content from {tables_with_content}/{table_count} tables, "
                f"processed {rows_processed} rows"
            )

            # Combine text
            result_text = "\n\n".join(text_parts)
            processing_time = time.time() - start_time
            logger.info(
                f"Simplified parsing complete in {processing_time:.2f}s, "
                f"generated {len(result_text)} characters of text"
            )

            # If the result is still empty, return an error message
            if not result_text:
                logger.warning("No text extracted using simplified method")
                return ""

            return result_text
        except Exception as backup_error:
            processing_time = time.time() - start_time
            logger.error(f"Simplified parsing failed after {processing_time:.2f}s: {str(backup_error)}")
            logger.error(f"Detailed traceback: {traceback.format_exc()}")
            return ""

class Docx:
    def __init__(self, max_image_size=1920, enable_multimodal=False, upload_file=None):
        logger.info("Initializing DOCX processor")
        self.max_image_size = max_image_size  # Maximum image size limit
        self.picture_cache = {}  # Image cache to avoid processing the same image repeatedly
        self.enable_multimodal = enable_multimodal
        self.upload_file = upload_file

    def get_picture(self, document, paragraph) -> Optional[Image.Image]:
        logger.info("Extracting image from paragraph")
        img = paragraph._element.xpath(".//pic:pic")
        if not img:
            logger.info("No image found in paragraph")
            return None
        img = img[0]
        try:
            embed = img.xpath(".//a:blip/@r:embed")[0]
            related_part = document.part.related_parts[embed]
            logger.info(f"Found embedded image with ID: {embed}")

            try:
                image_blob = related_part.image.blob
            except UnrecognizedImageError:
                logger.warning("Unrecognized image format. Skipping image.")
                return None
            except UnexpectedEndOfFileError:
                logger.warning(
                    "EOF was unexpectedly encountered while reading an image stream. Skipping image."
                )
                return None
            except InvalidImageStreamError:
                logger.warning(
                    "The recognized image stream appears to be corrupted. Skipping image."
                )
                return None

            try:
                logger.info("Converting image blob to PIL Image")
                image = Image.open(BytesIO(image_blob)).convert("RGBA")
                logger.info(
                    f"Successfully extracted image, size: {image.width}x{image.height}"
                )
                return image
            except Exception as e:
                logger.error(f"Failed to open image: {str(e)}")
                return None
        except Exception as e:
            logger.error(f"Error extracting image: {str(e)}")
            return None

    def __clean(self, line):
        logger.info(f"Cleaning text line of length: {len(line)}")
        line = re.sub(r"\u3000", " ", line).strip()
        return line

    def _identify_page_paragraph_mapping(self, max_page=100000):
        """Identify the paragraph range included on each page

        Args:
            max_page: Maximum number of pages to process

        Returns:
            dict: Mapping of page numbers to lists of paragraph indices
        """
        start_time = time.time()
        logger.info(f"Identifying page to paragraph mapping (max_page={max_page})")
        page_to_paragraphs = {}
        current_page = 0

        # Initialize page 0
        page_to_paragraphs[current_page] = []

        # Record the total number of paragraphs processed
        total_paragraphs = len(self.doc.paragraphs)
        logger.info(f"Total paragraphs to map: {total_paragraphs}")

        # Heuristic method: estimate the number of paragraphs per page
        # For large documents, using a heuristic can reduce XML parsing overhead
        if total_paragraphs > 1000:
            logger.info("Large document detected, using heuristic paragraph mapping")
            estimated_paras_per_page = 25  # Estimate approximately 25 paragraphs per page

            # Create an estimated page mapping
            for p_idx in range(total_paragraphs):
                est_page = p_idx // estimated_paras_per_page
                if est_page > max_page:
                    logger.info(f"Reached max page limit ({max_page}) at paragraph {p_idx}, stopping paragraph mapping")
                    break

                if est_page not in page_to_paragraphs:
                    page_to_paragraphs[est_page] = []

                page_to_paragraphs[est_page].append(p_idx)

                if p_idx > 0 and p_idx % 1000 == 0:
                    logger.info(f"Heuristic mapping: processed {p_idx}/{total_paragraphs} paragraphs")

            mapping_time = time.time() - start_time
            logger.info(f"Created heuristic mapping with {len(page_to_paragraphs)} pages in {mapping_time:.2f}s")
            return page_to_paragraphs

        # Standard method: iterate through all paragraphs to find page breaks
        logger.info("Using standard paragraph mapping method")
        page_breaks_found = 0
        for p_idx, p in enumerate(self.doc.paragraphs):
            # Add the current paragraph to the current page
            page_to_paragraphs[current_page].append(p_idx)

            # Log every 100 paragraphs
            if p_idx > 0 and p_idx % 100 == 0:
                logger.info(f"Processed {p_idx}/{total_paragraphs} paragraphs in page mapping")

            # Check for page breaks
            page_break_found = False

            # Method 1: Check for lastRenderedPageBreak
            for run in p.runs:
                if "lastRenderedPageBreak" in run._element.xml:
                    page_break_found = True
                    break

                if "w:br" in run._element.xml and 'type="page"' in run._element.xml:
                    page_break_found = True
                    break

            # Method 2: Check sectPr element (section break, usually indicates a new page)
            if not page_break_found and p._element.xpath(".//w:sectPr"):
                page_break_found = True

            # If a page break is found, create a new page
            if page_break_found:
                page_breaks_found += 1
                current_page += 1
                if current_page > max_page:
                    logger.info(f"Reached max page limit ({max_page}), stopping page mapping")
                    break

                # Initialize the paragraph list for the new page
                if current_page not in page_to_paragraphs:
                    page_to_paragraphs[current_page] = []

                if page_breaks_found % 10 == 0:
                    logger.info(f"Found {page_breaks_found} page breaks so far, current page: {current_page}")

        # Handle potential empty page mappings
        empty_pages = [page for page, paras in page_to_paragraphs.items() if not paras]
        if empty_pages:
            logger.info(f"Removing {len(empty_pages)} empty pages from mapping")
            for page in empty_pages:
                del page_to_paragraphs[page]

        mapping_time = time.time() - start_time
        logger.info(f"Created paragraph mapping with {len(page_to_paragraphs)} pages in {mapping_time:.2f}s")

        # Check the validity of the result
        if not page_to_paragraphs:
            logger.warning("No valid page mapping created, using fallback method")
            # All paragraphs are on page 0
            page_to_paragraphs[0] = list(range(total_paragraphs))

        # Log page distribution statistics
        page_sizes = [len(paragraphs) for paragraphs in page_to_paragraphs.values()]
        if page_sizes:
            avg_paragraphs = sum(page_sizes) / len(page_sizes)
            min_paragraphs = min(page_sizes)
            max_paragraphs = max(page_sizes)
            logger.info(
                f"Page statistics: avg={avg_paragraphs:.1f}, "
                f"min={min_paragraphs}, max={max_paragraphs} paragraphs per page"
            )

        return page_to_paragraphs


    def _process_page(self, page_num, paragraphs, from_page, to_page):
        """Process the content of one page

        Args:
            page_num: Page number
            paragraphs: List of paragraph indices
            from_page: Starting page number
            to_page: Ending page number

        Returns:
            list: List of processed result lines
        """
        logger.info(f"Processing page {page_num} with {len(paragraphs)} paragraphs")

        # Initialize thread local storage
        if not hasattr(thread_local, 'last_text'):
            thread_local.last_text = ""
            thread_local.last_image = None

        # Track content in sequence
        content_sequence = []
        current_text = ""
        current_images = []

        for para_idx in paragraphs:
            if para_idx >= len(self.doc.paragraphs):
                logger.warning(f"Paragraph index {para_idx} out of range")
                continue

            paragraph = self.doc.paragraphs[para_idx]

            # Get paragraph text
            text = paragraph.text.strip()
            if text:
                cleaned_text = self.__clean(text)
                
                # If we have pending images, add text after them
                if current_images:
                    # Add all accumulated images to the sequence
                    for img in current_images:
                        content_sequence.append(("image", img))
                    current_images = []
                    
                # Accumulate text
                current_text += cleaned_text + "\n"
            
            # Attempt to extract image
            image = self.get_picture(self.doc, paragraph)
            if image:
                # If we have text, add it to sequence first
                if current_text:
                    content_sequence.append(("text", current_text))
                    current_text = ""
                
                # Add image directly to sequence to maintain order
                content_sequence.append(("image", image))

        # Add any remaining text
        if current_text:
            content_sequence.append(("text", current_text))

        # Add any remaining images
        for img in current_images:
            content_sequence.append(("image", img))

        # Extract text and images in their original sequence for compatibility
        text_parts = []
        images = []
        
        for content_type, content in content_sequence:
            if content_type == "text":
                text_parts.append(content)
            else:  # image
                images.append(content)
        
        combined_text = "\n\n".join(text_parts) if text_parts else ""

        # After processing all paragraphs, add to results
        line_data = (combined_text, images, "", page_num, content_sequence)

        # Add to page-specific results
        page_lines = [line_data]

        # Add to global results
        with self.lines_lock:
            self.all_lines.append(line_data)

        # Update thread local storage
        thread_local.last_text = combined_text
        thread_local.last_image = images[-1] if images else None

        # Set last_image for this page
        if images:
            with self.last_image_lock:
                self.last_image_map[page_num] = images[-1]

        return page_lines

    def __call__(self, binary=None, from_page=0, to_page=100000, max_workers=None):
        """
        Process DOCX document, supporting concurrent processing of each page

        Args:
            binary: DOCX document binary content
            from_page: Starting page number
            to_page: Ending page number
            max_workers: Maximum number of workers, default to None (system decides)

        Returns:
            tuple: (List of text lines, List of tables)
        """
        logger.info("Processing DOCX document")

        # Check CPU core count to determine parallel strategy
        cpu_count = os.cpu_count() or 2
        logger.info(f"System has {cpu_count} CPU cores available")

        # Load document
        self.doc = self._load_document(binary)
        if not self.doc:
            return [], []

        # Identify page structure
        self.para_page_mapping = self._identify_page_paragraph_mapping(to_page)
        logger.info(f"Identified page to paragraph mapping for {len(self.para_page_mapping)} pages")

        # Apply page limits
        pages_to_process = self._apply_page_limit(self.para_page_mapping, from_page, to_page)
        if not pages_to_process:
            logger.warning("No pages to process after applying page limits!")
            return [], []

        # Initialize shared resources
        self._init_shared_resources()

        # Get current request ID
        current_request_id = self._get_request_id()

        # Process document content
        doc_size_mb = sys.getsizeof(binary) / (1024 * 1024)
        logger.info(f"Document binary size: {doc_size_mb:.2f} MB")

        # Decide on processing strategy
        if len(self.doc.paragraphs) < 30:
            # For small documents, use single-threaded processing
            self._process_small_document(pages_to_process, from_page, to_page)
        else:
            # For large documents, use multiprocessing
            self._process_large_document(binary,
                                         pages_to_process,
                                         from_page,
                                         to_page,
                                         max_workers,
                                         doc_size_mb,
                                         current_request_id,
                                         )

        # Process tables
        tbls = self._process_tables()

        # Clean up document resources
        self.doc = None
        self.last_image_map = None

        logger.info(f"Document processing complete, "
                    f"extracted {len(self.all_lines)} text sections and {len(tbls)} tables")
        return self.all_lines, tbls

    def _load_document(self, binary):
        """Load document

        Args:
            binary: Document binary content

        Returns:
            Document: Document object, or None (if loading fails)
        """
        try:
            doc = Document(BytesIO(binary))
            logger.info("Successfully loaded document from binary content")
            return doc
        except Exception as e:
            logger.error(f"Failed to load DOCX document: {str(e)}")
            return None

    def _init_shared_resources(self):
        """Initialize shared resources"""
        # Create shared resource locks to protect data structures shared between threads
        self.lines_lock = threading.Lock()
        self.last_image_lock = threading.Lock()

        # Initialize result containers
        self.all_lines = []
        self.last_image_map = {}  # Last image for each page

    def _get_request_id(self):
        """Get current request ID"""
        current_request_id = None
        try:
            from utils.request import get_request_id
            current_request_id = get_request_id()
            logger.info(f"Getting current request ID: {current_request_id} to pass to processing threads")
        except Exception as e:
            logger.warning(f"Failed to get current request ID: {str(e)}")
        return current_request_id

    def _apply_page_limit(self, para_page_mapping, from_page, to_page):
        """Apply page limits, return the list of pages to process

        Args:
            para_page_mapping: Mapping of pages to paragraphs
            from_page: Starting page number
            to_page: Ending page number

        Returns:
            list: List of pages to process
        """
        # Add page limits
        total_pages = len(para_page_mapping)
        if total_pages > to_page:
            logger.info(f"Document has {total_pages} pages, limiting processing to first {to_page} pages")
            logger.info(f"Setting to_page limit to {to_page}")
        else:
            logger.info(f"Document has {total_pages} pages, processing all pages (limit: {to_page})")

        # Filter out pages outside the range
        all_pages = sorted(para_page_mapping.keys())
        pages_to_process = [p for p in all_pages if from_page <= p < to_page]

        # Output the actual number of pages processed for debugging
        if pages_to_process:
            logger.info(
                f"Will process {len(pages_to_process)} pages "
                f"from page {from_page} to page {min(to_page, pages_to_process[-1] if pages_to_process else from_page)}"
            )

            if len(pages_to_process) < len(all_pages):
                logger.info(f"Skipping {len(all_pages) - len(pages_to_process)} pages due to page limit")

            # Log detailed page index information
            if len(pages_to_process) <= 10:
                logger.info(f"Pages to process: {pages_to_process}")
            else:
                logger.info(f"First 5 pages to process: {pages_to_process[:5]}, last 5: {pages_to_process[-5:]}")

        return pages_to_process

    def _process_small_document(self, pages_to_process, from_page, to_page):
        """Process small documents (less than 30 paragraphs)

        Args:
            pages_to_process: List of pages to process
            from_page: Starting page number
            to_page: Ending page number
        """
        logger.info(f"Small document detected ({len(self.doc.paragraphs)} paragraphs), "
                    f"processing without threading")
        for page_num in sorted(pages_to_process):
            if from_page <= page_num < to_page:
                logger.info(f"Processing page {page_num} (single-threaded mode)")
                self._process_page(page_num, self.para_page_mapping[page_num], from_page, to_page)
                logger.info(f"Completed processing page {page_num} (single-threaded mode)")
            else:
                logger.info(f"Skipping page {page_num} (out of requested range: {from_page}-{to_page})")

    def _process_large_document(self,
                                binary,
                                pages_to_process,
                                from_page,
                                to_page,
                                max_workers,
                                doc_size_mb,
                                current_request_id,
                                ):
        """Process large documents, using multiprocessing

        Args:
            binary: Document binary content
            pages_to_process: List of pages to process
            from_page: Starting page number
            to_page: Ending page number
            max_workers: Maximum number of workers
            doc_size_mb: Document size (MB)
            current_request_id: Current request ID
        """
        # If the number of pages is too large, process in batches to reduce memory consumption
        cpu_count = os.cpu_count() or 2

        # Check if the document contains images to optimize processing speed
        doc_contains_images = self._check_document_has_images()

        # Optimize process count: dynamically adjust based on number of pages and CPU cores
        if max_workers is None:
            max_workers = self._calculate_optimal_workers(doc_contains_images, pages_to_process, cpu_count)

        # Use temporary file to share large documents
        temp_file_path = self._prepare_document_sharing(binary, doc_size_mb)

        # Prepare multiprocess processing arguments
        args_list = self._prepare_multiprocess_args(
            pages_to_process, from_page, to_page, doc_contains_images,
            binary, temp_file_path
        )

        # Execute multiprocess tasks
        self._execute_multiprocess_tasks(args_list, max_workers)

        # Clean up temporary file
        self._cleanup_temp_file(temp_file_path)

    def _check_document_has_images(self):
        """Check if the document contains images

        Returns:
            bool: Whether the document contains images
        """
        doc_contains_images = False
        if hasattr(self.doc, 'inline_shapes') and len(self.doc.inline_shapes) > 0:
            doc_contains_images = True
            logger.info(f"Document contains {len(self.doc.inline_shapes)} inline images")
        return doc_contains_images

    def _calculate_optimal_workers(self, doc_contains_images, pages_to_process, cpu_count):
        """Calculate the optimal number of workers

        Args:
            doc_contains_images: Whether the document contains images
            pages_to_process: List of pages to process
            cpu_count: Number of CPU cores

        Returns:
            int: Optimal number of workers
        """
        # If no images or few pages, use fewer processes to avoid overhead
        if not doc_contains_images or len(pages_to_process) < cpu_count:
            max_workers = min(len(pages_to_process), max(1, cpu_count - 1))
        else:
            max_workers = min(len(pages_to_process), cpu_count)
        logger.info(f"Automatically set worker count to {max_workers}")
        return max_workers

    def _prepare_document_sharing(self, binary, doc_size_mb):
        """Prepare document sharing method

        Args:
            binary: Document binary content
            doc_size_mb: Document size (MB)

        Returns:
            str: Temporary file path, or None if not using
        """
        # For large documents, consider using temporary file to share data
        use_temp_file = doc_size_mb > 50  # If document is larger than 50MB, use temporary file for sharing
        temp_file_path = None

        if use_temp_file:
            logger.info(f"Large document detected ({doc_size_mb:.2f} MB), using temporary file for sharing")
            import tempfile
            temp_file = tempfile.NamedTemporaryFile(delete=False)
            temp_file_path = temp_file.name
            temp_file.write(binary)
            temp_file.close()
            logger.info(f"Wrote document to temporary file: {temp_file_path}")

        return temp_file_path

    def _prepare_multiprocess_args(self, pages_to_process, from_page, to_page,
                                 doc_contains_images, binary, temp_file_path):
        """Prepare a list of arguments for multiprocess processing

        Args:
            pages_to_process: List of pages to process
            from_page: Starting page number
            to_page: Ending page number
            doc_contains_images: Whether the document contains images
            binary: Document binary content
            temp_file_path: Temporary file path

        Returns:
            list: List of arguments
        """
        use_temp_file = temp_file_path is not None

        logger.info(f"Preparing multiprocess args with enable_multimodal={self.enable_multimodal}")

        # Pass parameters required by the page processing function
        args_list = []
        for page_num in pages_to_process:
            args_list.append((
                page_num,
                self.para_page_mapping[page_num],
                from_page,
                to_page,
                doc_contains_images,
                self.max_image_size,
                binary if not use_temp_file else None,
                temp_file_path if use_temp_file else None,
                self.enable_multimodal
            ))

        return args_list

    def _execute_multiprocess_tasks(self, args_list, max_workers):
        """Execute multiprocess tasks

        Args:
            args_list: List of arguments
            max_workers: Maximum number of workers
        """
        # Use a shared manager to share data
        with Manager() as manager:
            # Create shared data structures
            self.all_lines = manager.list()
            self.last_image_map = manager.dict()

            logger.info(f"Processing {len(args_list)} pages using {max_workers} processes")

            # Use ProcessPoolExecutor to truly implement multi-core parallelization
            batch_start_time = time.time()
            with ProcessPoolExecutor(max_workers=max_workers) as executor:
                logger.info(f"Started ProcessPoolExecutor with {max_workers} workers")

                # Submit all tasks
                future_to_idx = {
                    executor.submit(
                        process_page_multiprocess,
                        *args
                    ): i
                    for i, args in enumerate(args_list)
                }
                logger.info(f"Submitted {len(future_to_idx)} processing tasks to process pool")

                # Collect results
                self._collect_process_results(future_to_idx, args_list, batch_start_time)

    def _collect_process_results(self, future_to_idx, args_list, batch_start_time):
        """Collect multiprocess processing results

        Args:
            future_to_idx: Mapping of Future to index
            args_list: List of arguments
            batch_start_time: Batch start time
        """
        # Collect results
        completed_count = 0
        results = []
        temp_img_paths = set()  # Collect all temporary image paths

        for future in as_completed(future_to_idx):
            idx = future_to_idx[future]
            page_num = args_list[idx][0]
            try:
                page_lines = future.result()

                # Collect temporary image paths for later cleanup
                for line in page_lines:
                    if len(line) > 1 and isinstance(line[1], list):
                        for item in line[1]:
                            if isinstance(item, str) and item.startswith("/tmp/docx_img_"):
                                temp_img_paths.add(item)

                results.extend(page_lines)
                completed_count += 1

                if completed_count % max(1, len(args_list) // 10) == 0 or completed_count == len(args_list):
                    elapsed_ms = int((time.time() - batch_start_time) * 1000)
                    progress_pct = int((completed_count / len(args_list)) * 100)
                    logger.info(
                        f"Progress: {completed_count}/{len(args_list)} pages processed "
                        f"({progress_pct}%, elapsed: {elapsed_ms}ms)"
                    )

            except Exception as e:
                logger.error(f"Error processing page {page_num}: {str(e)}")
                logger.error(f"Detailed traceback for page {page_num}: {traceback.format_exc()}")

        # Process completion
        processing_elapsed_ms = int((time.time() - batch_start_time) * 1000)
        logger.info(f"All processing completed in {processing_elapsed_ms}ms")

        # Process results
        self._process_multiprocess_results(results)

        # Clean up temporary image files
        self._cleanup_temp_image_files(temp_img_paths)

    def _process_multiprocess_results(self, results):
        """Process multiprocess results

        Args:
            results: List of processed results
        """
        lines = list(results)

        # Process images - must be handled in the main process for upload
        # If images are being processed, they need to be handled in the main process for upload
        image_upload_start = time.time()
        
        # Count total images to process
        images_to_process = []
        for i, line_data in enumerate(lines):
            if len(line_data) > 1 and line_data[1]:  # Check if there are images
                # Confirm image data is valid - it might be lost during inter-process communication
                if isinstance(line_data[1], list) and len(line_data[1]) > 0:
                    images_to_process.append(i)
                    logger.info(f"Found line {i} with {len(line_data[1])} images to process")
                else:
                    logger.warning(f"Line {i} has invalid image data: {type(line_data[1])}")

        # Process images if needed
        image_url_map = {}  # Map from image path to URL
        if images_to_process:
            logger.info(f"Found {len(images_to_process)} lines with images to process in main process")
            
            # First, create a mapping of image paths to uploaded URLs
            for line_idx in images_to_process:
                line_data = lines[line_idx]
                _, image_paths, _, page_num = line_data[:4]
                
                # Process all image file paths
                for img_path in image_paths:
                    if os.path.exists(img_path) and img_path not in image_url_map:
                        try:
                            image_url = self.upload_file(img_path)
                            if image_url:
                                # Add image URL as Markdown format
                                markdown_image = f"![]({image_url})"
                                image_url_map[img_path] = markdown_image
                                logger.info(f"Added image URL for {img_path}: {image_url}")
                            else:
                                logger.warning(f"Failed to upload image: {img_path}")
                        except Exception as e:
                            logger.error(f"Error processing image from page {page_num}: {str(e)}")

            image_upload_elapsed = time.time() - image_upload_start
            logger.info(f"Finished uploading {len(image_url_map)} images in {image_upload_elapsed:.2f}s")
        
        # Process content in original sequence order
        processed_lines = []
        for line_data in lines:
            if len(line_data) >= 5 and line_data[4]:  # Check if we have processed_content
                processed_content = line_data[4]
                page_num = line_data[3]
                
                # Reconstruct text with images in original positions
                combined_parts = []
                for content_type, content in processed_content:
                    if content_type == "text":
                        combined_parts.append(content)
                    elif content_type == "image" and content in image_url_map:
                        combined_parts.append(image_url_map[content])
                
                # Create the final text with proper ordering
                final_text = "\n\n".join(part for part in combined_parts if part)
                processed_lines.append((final_text, None, "", page_num))
                
            else:
                # Fallback to original approach if no processed_content
                text, image_paths, _, page_num = line_data[:4]
                
                if image_paths:
                    # Process all image URLs
                    processed_urls = []
                    for img_path in image_paths:
                        if img_path in image_url_map:
                            processed_urls.append(image_url_map[img_path])
                    
                    if processed_urls:
                        # Add all images at the end (legacy approach)
                        image_text = "\n\n".join(processed_urls)
                        if text:
                            updated_text = text + "\n\n" + image_text
                        else:
                            updated_text = image_text
                        processed_lines.append((updated_text, None, "", page_num))
                    else:
                        processed_lines.append((text, None, "", page_num))
                else:
                    processed_lines.append((text, None, "", page_num))

        # Sort results by page number
        sorted_lines = sorted(processed_lines, key=lambda x: x[3] if len(x) > 3 else 0)
        self.all_lines = [(line[0], line[1], line[2]) for line in sorted_lines if len(line) > 2]
        
        logger.info(f"Finished processing {len(self.all_lines)} lines with interleaved images and text")

    def _cleanup_temp_image_files(self, temp_paths):
        """Clean up temporary image files created by multiprocessing

        Args:
            temp_paths: Set of temporary file paths
        """
        if not temp_paths:
            return

        logger.info(f"Cleaning up {len(temp_paths)} temporary image files")
        deleted_count = 0
        error_count = 0

        for path in temp_paths:
            try:
                if os.path.exists(path):
                    os.unlink(path)
                    deleted_count += 1
                    # Delete temporary directory (if empty)
                    try:
                        temp_dir = os.path.dirname(path)
                        if temp_dir.startswith("/tmp/docx_img_") and os.path.exists(temp_dir):
                            os.rmdir(temp_dir)
                    except OSError:
                        # If directory is not empty, ignore error
                        pass
            except Exception as e:
                logger.error(f"Failed to delete temp file {path}: {str(e)}")
                error_count += 1

        logger.info(f"Temporary file cleanup: deleted {deleted_count}, errors {error_count}")

    def _process_images_in_main_process(self, lines, images_to_process):
        """Process image uploads in the main process

        Args:
            lines: List of line data
            images_to_process: Indices of lines with images to process
        """
        # Process all images
        with tempfile.TemporaryDirectory() as temp_dir:
            for line_idx in images_to_process:
                line_data = lines[line_idx]
                text, image_data, _, page_num = line_data

                image_paths = []

                # If it's a list of image objects (single-process), save them as temporary files
                if image_data and isinstance(image_data, list):
                    if all(isinstance(item, str) for item in image_data):
                        # Multi-process mode, images are already saved as temporary files
                        logger.info(f"Found {len(image_data)} image paths for page {page_num}")
                        image_paths = image_data
                    elif all(hasattr(item, 'save') for item in image_data):
                        # Single-process mode, need to save image objects as temporary files
                        logger.info(f"Converting {len(image_data)} image objects to files for page {page_num}")
                        for img_idx, image in enumerate(image_data):
                            try:
                                # Save temporary image file
                                temp_file_path = os.path.join(temp_dir, f"page_{page_num}_img_{img_idx}.png")
                                image.save(temp_file_path, format="PNG")
                                image_paths.append(temp_file_path)
                                logger.info(f"Saved image from page {page_num} to {temp_file_path}")
                            except Exception as e:
                                logger.error(f"Failed to save image: {str(e)}")
                    else:
                        logger.warning(f"Unknown image data format for page {page_num}: {type(image_data[0])}")

                # Process all image file paths
                processed_urls = []
                for img_path in image_paths:
                    if os.path.exists(img_path):
                        try:
                            image_url = self.upload_file(img_path)
                            if image_url:
                                # Add image URL as Markdown format
                                markdown_image = f"![]({image_url})"
                                processed_urls.append(markdown_image)
                                logger.info(f"Added image URL to page {page_num} result: {image_url}")
                            else:
                                logger.warning(f"Failed to upload image: {img_path}")

                            # Do not delete temporary files created by multiprocessing, they will be cleaned up later
                            if not img_path.startswith("/tmp/docx_img_"):
                                try:
                                    os.unlink(img_path)
                                    logger.info(f"Deleted temporary file: {img_path}")
                                except Exception as e:
                                    logger.error(f"Failed to delete temporary file: {str(e)}")

                        except Exception as e:
                            logger.error(f"Error processing image from page {page_num}: {str(e)}")

                # Update text, add all image links
                if processed_urls:
                    image_text = "\n\n".join(processed_urls)
                    if text:
                        updated_text = text + "\n\n" + image_text
                    else:
                        updated_text = image_text

                    # Update text in results, clear image list (since processing is complete)
                    lines[line_idx] = (updated_text, [], "", page_num)
                    logger.info(f"Updated text for page {page_num} with {len(processed_urls)} image URLs")

    def _cleanup_temp_file(self, temp_file_path):
        """Clean up temporary file

        Args:
            temp_file_path: Temporary file path
        """
        if temp_file_path and os.path.exists(temp_file_path):
            try:
                os.unlink(temp_file_path)
                logger.info(f"Removed temporary file: {temp_file_path}")
            except Exception as e:
                logger.error(f"Failed to remove temporary file: {str(e)}")

    def _process_tables(self):
        """Process tables in the document

        Returns:
            list: List of tables
        """
        tbls = []
        table_count = len(self.doc.tables)
        if table_count > 0:
            logger.info(f"Processing {table_count} tables")
            for tb_idx, tb in enumerate(self.doc.tables):
                if tb_idx % 10 == 0:  # Log only every 10 tables to reduce log volume
                    logger.info(f"Processing table {tb_idx+1}/{table_count}")

                # Optimize: Check if table is empty
                if len(tb.rows) == 0 or all(len(r.cells) == 0 for r in tb.rows):
                    logger.info(f"Skipping empty table {tb_idx+1}")
                    continue

                table_html = self._convert_table_to_html(tb)
                tbls.append(((None, table_html), ""))

        return tbls

    def _convert_table_to_html(self, table):
        """Convert table to HTML

        Args:
            table: Table object

        Returns:
            str: HTML formatted table
        """
        html = "<table>"
        for r in table.rows:
            html += "<tr>"
            i = 0
            while i < len(r.cells):
                span = 1
                c = r.cells[i]
                for j in range(i + 1, len(r.cells)):
                    if c.text == r.cells[j].text:
                        span += 1
                        i = j
                i += 1
                html += (
                    f"<td>{c.text}</td>"
                    if span == 1
                    else f"<td colspan='{span}'>{c.text}</td>"
                )
            html += "</tr>"
        html += "</table>"
        return html

    def _safe_concat_images(self, images):
        """Safely concatenate image lists

        Args:
            images: List of images

        Returns:
            Image: Concatenated image, or the first image (if concatenation fails)
        """
        if not images:
            return None

        if len(images) == 1:
            return images[0]

        try:
            logger.info(f"Attempting to concatenate {len(images)} images")
            from PIL import Image

            # Calculate the size of the concatenated image
            total_width = max(img.width for img in images if hasattr(img, 'width'))
            total_height = sum(img.height for img in images if hasattr(img, 'height'))

            if total_width <= 0 or total_height <= 0:
                logger.warning("Invalid image size, returning the first image")
                return images[0]

            # Create a new image
            new_image = Image.new("RGBA", (total_width, total_height), (0, 0, 0, 0))

            # Paste images one by one
            y_offset = 0
            for img in images:
                if not hasattr(img, 'width') or not hasattr(img, 'height'):
                    continue

                new_image.paste(img, (0, y_offset))
                y_offset += img.height

            logger.info(f"Successfully concatenated images, final size: {total_width}x{total_height}")
            return new_image
        except Exception as e:
            logger.error(f"Failed to concatenate images: {str(e)}")
            logger.error(f"Detailed error: {traceback.format_exc()}")
            # If concatenation fails, return the first image
            return images[0]



def _save_image_to_temp(logger, image, page_num, img_idx):
    """Save image to a temporary file to pass between processes

    Args:
        logger: Logger
        image: PIL image object
        page_num: Page number
        img_idx: Image index

    Returns:
        str: Temporary file path, or None (if saving fails)
    """
    if not image:
        return None

    import tempfile
    import os

    try:
        # Create a temporary file
        temp_dir = tempfile.mkdtemp(prefix="docx_img_")
        temp_file_path = os.path.join(temp_dir, f"page_{page_num}_img_{img_idx}.png")

        # Save the image
        image.save(temp_file_path, format="PNG")
        logger.info(f"[PID:{os.getpid()}] Saved image to temporary file: {temp_file_path}")

        return temp_file_path
    except Exception as e:
        logger.error(f"[PID:{os.getpid()}] Failed to save image to temp file: {str(e)}")
        return None


def process_page_multiprocess(page_num,
                              paragraphs,
                              from_page,
                              to_page,
                              doc_contains_images,
                              max_image_size,
                              doc_binary,
                              temp_file_path,
                              enable_multimodal,
                              ):
    """Page processing function specifically designed for multiprocessing

    Args:
        page_num: Page number
        paragraphs: List of paragraph indices
        from_page: Starting page number
        to_page: Ending page number
        doc_contains_images: Whether the document contains images
        max_image_size: Maximum image size
        doc_binary: Document binary content
        temp_file_path: Temporary file path, if using
        enable_multimodal: Whether to enable multimodal processing

    Returns:
        list: List of processed result lines
    """
    try:
        # Set process-level logging
        process_logger = logging.getLogger(__name__)

        # If outside processing range, do not process
        if page_num < from_page or page_num >= to_page:
            process_logger.info(f"[PID:{os.getpid()}] Skipping page {page_num} (out of requested range)")
            return []

        process_logger.info(f"[PID:{os.getpid()}] Processing page {page_num} with {len(paragraphs)} paragraphs, "
                             f"enable_multimodal={enable_multimodal}")
        start_time = time.time()

        # Load document in the process
        doc = _load_document_in_process(process_logger, page_num, doc_binary, temp_file_path)
        if not doc:
            return []

        # If paragraph indices are empty, return empty result
        if not paragraphs:
            process_logger.info(f"[PID:{os.getpid()}] No paragraphs to process for page {page_num}")
            return []

        # Extract page content
        combined_text, images, content_sequence = _extract_page_content_in_process(
            process_logger,
            doc,
            page_num,
            paragraphs,
            enable_multimodal,
            max_image_size
        )

        # Process content sequence to maintain order between processes
        processed_content = []
        temp_image_index = 0
        
        if enable_multimodal:
            # First pass: save all images to temporary files
            image_paths = []
            for i, img in enumerate(images):
                img_path = _save_image_to_temp(process_logger, img, page_num, i)
                if img_path:
                    image_paths.append(img_path)
                    
            process_logger.info(
                f"[PID:{os.getpid()}] Saved {len(image_paths)} images to temp files for page {page_num}"
            )
            
            # Second pass: reconstruct the content sequence with image paths
            for content_type, content in content_sequence:
                if content_type == "text":
                    processed_content.append(("text", content))
                else:  # image
                    if temp_image_index < len(image_paths):
                        processed_content.append(("image", image_paths[temp_image_index]))
                        temp_image_index += 1

        # Create result line with the ordered content sequence
        line_data = (combined_text, image_paths, "", page_num, processed_content)
        page_lines = [line_data]

        processing_time = time.time() - start_time
        process_logger.info(
            f"[PID:{os.getpid()}] Page {page_num} processing completed in {processing_time:.2f}s"
        )

        return page_lines

    except Exception as e:
        process_logger = logging.getLogger(__name__)
        process_logger.error(f"[PID:{os.getpid()}] Error processing page {page_num}: {str(e)}")
        process_logger.error(f"[PID:{os.getpid()}] Traceback: {traceback.format_exc()}")
        return []


def _load_document_in_process(logger, page_num, doc_binary, temp_file_path):
    """Load document in a process

    Args:
        logger: Logger
        page_num: Page number
        doc_binary: Document binary data
        temp_file_path: Temporary file path

    Returns:
        Document: Loaded document object, or None (if loading fails)
    """
    logger.info(f"[PID:{os.getpid()}] Loading document in process for page {page_num}")
    try:
        from docx import Document
        from io import BytesIO

        # Load document based on the provided method
        if doc_binary is not None:
            # Load document from binary content
            doc = Document(BytesIO(doc_binary))
            logger.info(f"[PID:{os.getpid()}] Loaded document from binary data")
        elif temp_file_path is not None:
            # Load document from temporary file
            doc = Document(temp_file_path)
            logger.info(f"[PID:{os.getpid()}] Loaded document from temp file: {temp_file_path}")
        else:
            logger.error(f"[PID:{os.getpid()}] No document source provided")
            return None

        logger.info(f"[PID:{os.getpid()}] "
            f"Successfully loaded document with {len(doc.paragraphs)} paragraphs")
        return doc

    except Exception as e:
        logger.error(f"[PID:{os.getpid()}] Failed to load document: {str(e)}")
        logger.error(f"[PID:{os.getpid()}] Error traceback: {traceback.format_exc()}")
        return None


def _extract_page_content_in_process(logger, doc, page_num, paragraphs, enable_multimodal, max_image_size):
    """Extract page content in a process

    Args:
        logger: Logger
        doc: Document object
        page_num: Page number
        paragraphs: List of paragraph indices
        enable_multimodal: Whether to enable multimodal processing
        max_image_size: Maximum image size

    Returns:
        tuple: (Extracted text, List of extracted images)
    """
    logger.info(f"[PID:{os.getpid()}] Page {page_num}: Processing {len(paragraphs)} paragraphs, "
                f"enable_multimodal={enable_multimodal}")

    # Instead of separate collections, track content in paragraph sequence
    content_sequence = []
    current_text = ""
    
    processed_paragraphs = 0
    paragraphs_with_text = 0
    paragraphs_with_images = 0

    for para_idx in paragraphs:
        if para_idx >= len(doc.paragraphs):
            logger.warning(f"[PID:{os.getpid()}] Paragraph index {para_idx} out of range")
            continue

        paragraph = doc.paragraphs[para_idx]
        processed_paragraphs += 1
        

        # Extract text content
        text = paragraph.text.strip()
        if text:
            # Clean text
            cleaned_text = re.sub(r"\u3000", " ", text).strip()
            current_text += cleaned_text + "\n"
            paragraphs_with_text += 1

        # Process image - if multimodal processing is enabled
        if enable_multimodal:
            image = _extract_image_in_process(logger, doc, paragraph, page_num, para_idx, max_image_size)
            if image:
                # If we have accumulated text, add it to sequence first
                if current_text:
                    content_sequence.append(("text", current_text))
                    current_text = ""
                
                # Add image to sequence
                content_sequence.append(("image", image))
                paragraphs_with_images += 1

        if processed_paragraphs % 50 == 0:
            logger.info(
                f"[PID:{os.getpid()}] "
                f"Page {page_num}: Processed {processed_paragraphs}/{len(paragraphs)} paragraphs"
            )

    # Add any remaining text
    if current_text:
        content_sequence.append(("text", current_text))

    logger.info(
        f"[PID:{os.getpid()}] Page {page_num}: Completed content extraction, "
        f"found {paragraphs_with_text} paragraphs with text, "
        f"{paragraphs_with_images} with images, "
        f"total content items: {len(content_sequence)}"
    )

    # Extract text and images in their original sequence
    text_parts = []
    images = []
    
    for content_type, content in content_sequence:
        if content_type == "text":
            text_parts.append(content)
        else:  # image
            images.append(content)
    
    combined_text = "\n\n".join(text_parts) if text_parts else ""
    
    return combined_text, images, content_sequence


def _extract_image_in_process(logger, doc, paragraph, page_num, para_idx, max_image_size):
    """Extract image from a paragraph in a process

    Args:
        logger: Logger
        doc: Document object
        paragraph: Paragraph object
        page_num: Page number
        para_idx: Paragraph index
        max_image_size: Maximum image size

    Returns:
        Image: Extracted image object, or None
    """
    try:
        # Attempt to extract image
        img = paragraph._element.xpath(".//pic:pic")
        if not img:
            logger.warning(f"[PID:{os.getpid()}] Page {page_num}: No pic element found in paragraph {para_idx}")
            return None

        img = img[0]
        logger.info(f"[PID:{os.getpid()}] Page {page_num}: Found pic element in paragraph {para_idx}")

        try:
            # Extract image ID and related part
            embed = img.xpath(".//a:blip/@r:embed")
            if not embed:
                logger.warning(f"[PID:{os.getpid()}] Page {page_num}: No embed attribute found in image")
                return None

            embed = embed[0]
            if embed not in doc.part.related_parts:
                logger.warning(f"[PID:{os.getpid()}] Page {page_num}: Embed ID {embed} not found in related parts")
                return None

            related_part = doc.part.related_parts[embed]
            logger.info(f"[PID:{os.getpid()}] Found embedded image with ID: {embed}")

            # Attempt to get image data
            try:
                image_blob = related_part.image.blob
                logger.info(f"[PID:{os.getpid()}] Successfully extracted image blob, size: {len(image_blob)} bytes")
            except Exception as blob_error:
                logger.warning(
                    f"[PID:{os.getpid()}] Error extracting image blob: {str(blob_error)}"
                )
                return None

            # Convert data to PIL image
            try:
                from PIL import Image
                from io import BytesIO
                image = Image.open(BytesIO(image_blob)).convert("RGBA")

                # Check image size
                if hasattr(image, 'width') and hasattr(image, 'height'):
                    logger.info(
                        f"[PID:{os.getpid()}] Successfully created image object, "
                        f"size: {image.width}x{image.height}"
                    )

                    # Skip small images (usually decorative elements)
                    if image.width < 50 or image.height < 50:
                        logger.info(
                            f"[PID:{os.getpid()}] "
                            f"Skipping small image ({image.width}x{image.height})"
                        )
                        return None

                    # Scale large images
                    if image.width > max_image_size or image.height > max_image_size:
                        scale = min(max_image_size / image.width, max_image_size / image.height)
                        new_width = int(image.width * scale)
                        new_height = int(image.height * scale)
                        resized_image = image.resize((new_width, new_height))
                        logger.info(f"[PID:{os.getpid()}] Resized image to {new_width}x{new_height}")
                        return resized_image

                logger.info(f"[PID:{os.getpid()}] Found image in paragraph {para_idx}")
                return image
            except Exception as e:
                logger.error(f"[PID:{os.getpid()}] Failed to create image from blob: {str(e)}")
                logger.error(f"[PID:{os.getpid()}] Error traceback: {traceback.format_exc()}")
                return None
        except Exception as e:
            logger.error(f"[PID:{os.getpid()}] Error extracting image: {str(e)}")
            logger.error(f"[PID:{os.getpid()}] Error traceback: {traceback.format_exc()}")
            return None
    except Exception as e:
        logger.error(f"[PID:{os.getpid()}] Error processing image: {str(e)}")
        logger.error(f"[PID:{os.getpid()}] Error traceback: {traceback.format_exc()}")
        return None
    finally:
        logger.info(f"[PID:{os.getpid()}] Page {page_num}: Image extraction completed")
